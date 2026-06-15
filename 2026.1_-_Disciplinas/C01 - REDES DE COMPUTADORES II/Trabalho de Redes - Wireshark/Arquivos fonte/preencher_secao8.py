# -*- coding: utf-8 -*-
"""
Script para preencher a Secao 8 do Relatorio Final.docx
com texto de analise, tabelas e imagens de graficos.
Resolve caminhos com acentos dinamicamente via glob.
"""
import os
import glob
import shutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================
# Resolver caminhos (lida com acentos)
# ============================================================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.dirname(SCRIPT_DIR)

# Encontrar o docx
docx_matches = glob.glob(os.path.join(BASE_DIR, "Documento final", "*.docx"))
if not docx_matches:
    print("ERRO: Nenhum arquivo .docx encontrado!")
    exit(1)
DOCX_PATH = docx_matches[0]
print(f"DOCX encontrado: {DOCX_PATH}")

# Encontrar pasta de estatisticas
stat_matches = glob.glob(os.path.join(BASE_DIR, "Prints", "5 -*"))
if not stat_matches:
    print("ERRO: Pasta de estatisticas nao encontrada!")
    exit(1)
IMG_DIR = stat_matches[0]
print(f"Pasta de imagens: {IMG_DIR}")

# Encontrar fig8
fig8_matches = glob.glob(os.path.join(BASE_DIR, "Prints", "4 - ICMP", "fig8_protocolos.png"))
FIG8_PATH = fig8_matches[0] if fig8_matches else None

# Backup
backup_path = DOCX_PATH.replace(".docx", "_BACKUP.docx")
if not os.path.exists(backup_path):
    shutil.copy2(DOCX_PATH, backup_path)
    print(f"Backup: {backup_path}")

# ============================================================
# Abrir documento
# ============================================================
doc = Document(DOCX_PATH)

# Encontrar secao 8 e 9
section8_idx = None
section9_idx = None

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if section8_idx is None and '8.' in text and 'stat' in text.lower() or (section8_idx is None and '8.' in text and 'Estat' in text):
        section8_idx = i
    elif section8_idx is not None and section9_idx is None and ('9.' in text and ('Discuss' in text or 'discuss' in text)):
        section9_idx = i
        break

if section8_idx is None:
    # Tentar busca mais flexivel
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if '8.' in text and 'Gr' in text:
            section8_idx = i
        elif section8_idx is not None and '9.' in text:
            section9_idx = i
            break

if section8_idx is None:
    print("ERRO: Secao 8 nao encontrada!")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"  [{i}] {para.text.strip()[:80]}")
    exit(1)

print(f"Secao 8: paragrafo {section8_idx} -> '{doc.paragraphs[section8_idx].text.strip()[:60]}'")
print(f"Secao 9: paragrafo {section9_idx} -> '{doc.paragraphs[section9_idx].text.strip()[:60]}'")

# Remover placeholders entre secao 8 e 9
for idx in reversed(range(section8_idx + 1, section9_idx)):
    p_element = doc.paragraphs[idx]._element
    p_element.getparent().remove(p_element)
print(f"Removidos {section9_idx - section8_idx - 1} paragrafos de placeholder")

# ============================================================
# Funcao para inserir conteudo apos referencia
# ============================================================
def insert_after(ref, new_p):
    ref._element.addnext(new_p._element)

def make_paragraph(text, font_size=12, bold=False, italic=False, color=None, alignment=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment:
        p.alignment = alignment
    return p

def make_image_paragraph(img_path, width_inches=5.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    return p

def make_caption(text):
    return make_paragraph(text, font_size=10, italic=True, color=(100, 100, 100), alignment=WD_ALIGN_PARAGRAPH.CENTER)

# ============================================================
# Construir conteudo da secao 8
# ============================================================
ref = doc.paragraphs[section8_idx]

# Espacamento e Introducao
elements = []

# Introducao
elements.append(make_paragraph(
    u"Para a an\u00e1lise estat\u00edstica da captura de tr\u00e1fego de rede, foi utilizada a funcionalidade "
    u"Protocol Hierarchy Statistics do Wireshark, acess\u00edvel pelo menu Statistics > Protocol Hierarchy. "
    u"Essa ferramenta permite visualizar todos os protocolos presentes na captura, organizados hierarquicamente "
    u"de acordo com o modelo de camadas, exibindo a quantidade de pacotes, a porcentagem relativa e o volume "
    u"de bytes de cada protocolo."
))

# 8.1 Hierarquia de Protocolos
elements.append(make_paragraph(u"8.1 Hierarquia de Protocolos", font_size=14, bold=True))
elements.append(make_paragraph(
    u"A Figura abaixo apresenta a tela do Protocol Hierarchy Statistics do Wireshark, exibindo todos os "
    u"protocolos identificados durante a captura de tr\u00e1fego. Foram capturados um total de 2.275 pacotes, "
    u"correspondendo a 1.244.568 bytes (aproximadamente 1,2 MB) de dados."
))

if FIG8_PATH and os.path.exists(FIG8_PATH):
    elements.append(make_image_paragraph(FIG8_PATH, 5.8))
    elements.append(make_caption(u"Figura 1 \u2013 Hierarquia de protocolos (Protocol Hierarchy Statistics \u2013 Wireshark)"))
    print("  Imagem Protocol Hierarchy inserida")

elements.append(make_paragraph(
    u"A an\u00e1lise da hierarquia de protocolos revelou que o tr\u00e1fego capturado \u00e9 predominantemente "
    u"composto por protocolos da pilha TCP/IP, com destaque para a presen\u00e7a significativa de tr\u00e1fego "
    u"IPv6 (74,7% dos pacotes) em compara\u00e7\u00e3o ao IPv4 (24,0%). Essa distribui\u00e7\u00e3o reflete a "
    u"tend\u00eancia de ado\u00e7\u00e3o crescente do IPv6 nas redes modernas."
))

# 8.2 Quantidade de Pacotes por Protocolo
elements.append(make_paragraph(u"8.2 Quantidade de Pacotes por Protocolo", font_size=14, bold=True))
elements.append(make_paragraph(
    u"A tabela a seguir apresenta a quantidade de pacotes e bytes organizados por protocolo de transporte, "
    u"consolidando os dados de IPv4 e IPv6:"
))

tabela_path = os.path.join(IMG_DIR, "tabela_protocolos.png")
if os.path.exists(tabela_path):
    elements.append(make_image_paragraph(tabela_path, 5.5))
    elements.append(make_caption(u"Tabela 1 \u2013 Quantidade de pacotes e bytes por protocolo de transporte"))
    print("  Tabela inserida")

elements.append(make_paragraph(
    u"Conforme demonstrado na tabela, o protocolo TCP concentra a maior parte do tr\u00e1fego, com 1.300 pacotes "
    u"(57,1% do total), seguido pelo UDP com 626 pacotes (27,5%). O ICMPv6 aparece em terceiro lugar com 230 "
    u"pacotes (10,1%), refletindo o uso de mensagens de controle e descoberta de vizinhos no IPv6."
))

# 8.3 Graficos
elements.append(make_paragraph(u"8.3 Gr\u00e1ficos de Distribui\u00e7\u00e3o", font_size=14, bold=True))
elements.append(make_paragraph(
    u"Para melhor visualiza\u00e7\u00e3o dos dados, foram elaborados gr\u00e1ficos que ilustram a "
    u"distribui\u00e7\u00e3o dos pacotes capturados:"
))

graf1_path = os.path.join(IMG_DIR, "grafico1_distribuicao_protocolos.png")
if os.path.exists(graf1_path):
    elements.append(make_image_paragraph(graf1_path, 5.2))
    elements.append(make_caption(u"Gr\u00e1fico 1 \u2013 Distribui\u00e7\u00e3o de pacotes por protocolo"))
    print("  Grafico 1 inserido")

elements.append(make_paragraph(
    u"O gr\u00e1fico acima evidencia a predomin\u00e2ncia do TCP no tr\u00e1fego capturado, o que \u00e9 "
    u"esperado considerando que a maior parte das aplica\u00e7\u00f5es web modernas utiliza conex\u00f5es TCP "
    u"para garantir a entrega confi\u00e1vel dos dados. O UDP, embora em menor quantidade de pacotes, desempenha "
    u"papel importante no transporte de protocolos como DNS e QUIC."
))

# 8.4 Trafego TCP
elements.append(make_paragraph(u"8.4 Tr\u00e1fego TCP", font_size=14, bold=True))
elements.append(make_paragraph(
    u"O protocolo TCP foi respons\u00e1vel por 1.300 pacotes (57,1% do total), distribu\u00eddos entre IPv6 "
    u"(845 pacotes) e IPv4 (455 pacotes). Dentro do tr\u00e1fego TCP, destaca-se a presen\u00e7a do protocolo "
    u"TLS (Transport Layer Security), com 440 pacotes no total (271 via IPv6 e 169 via IPv4), evidenciando o "
    u"uso predominante de comunica\u00e7\u00f5es criptografadas durante a captura."
))
elements.append(make_paragraph(
    u"Em termos de volume de dados, o TLS representou a maior parcela do tr\u00e1fego, com 650.793 bytes "
    u"(52,3% do total), confirmando que a maioria das comunica\u00e7\u00f5es capturadas utilizou criptografia "
    u"para proteger os dados transmitidos."
))

# 8.5 Trafego UDP
elements.append(make_paragraph(u"8.5 Tr\u00e1fego UDP", font_size=14, bold=True))
elements.append(make_paragraph(
    u"O protocolo UDP totalizou 626 pacotes (27,5% do total), com forte predomin\u00e2ncia no IPv6 (611 pacotes "
    u"contra apenas 15 no IPv4). O principal protocolo de aplica\u00e7\u00e3o identificado sobre UDP foi o QUIC "
    u"(442 pacotes), um protocolo moderno desenvolvido pelo Google que combina recursos de transporte e "
    u"criptografia, amplamente utilizado por navegadores e servi\u00e7os web atuais."
))
elements.append(make_paragraph(
    u"Tamb\u00e9m foram identificados 88 pacotes de DNS sobre UDP, consistente com as consultas de "
    u"resolu\u00e7\u00e3o de nomes realizadas durante a navega\u00e7\u00e3o."
))

graf2_path = os.path.join(IMG_DIR, "grafico2_tcp_vs_udp.png")
if os.path.exists(graf2_path):
    elements.append(make_image_paragraph(graf2_path, 4.8))
    elements.append(make_caption(u"Gr\u00e1fico 2 \u2013 Compara\u00e7\u00e3o de tr\u00e1fego TCP vs UDP (detalhamento por vers\u00e3o IP)"))
    print("  Grafico 2 inserido")

graf3_path = os.path.join(IMG_DIR, "grafico3_distribuicao_bytes.png")
if os.path.exists(graf3_path):
    elements.append(make_image_paragraph(graf3_path, 4.5))
    elements.append(make_caption(u"Gr\u00e1fico 3 \u2013 Distribui\u00e7\u00e3o do tr\u00e1fego por volume de bytes"))
    print("  Grafico 3 inserido")

graf5_path = os.path.join(IMG_DIR, "grafico5_protocolos_aplicacao.png")
if os.path.exists(graf5_path):
    elements.append(make_image_paragraph(graf5_path, 5.0))
    elements.append(make_caption(u"Gr\u00e1fico 4 \u2013 Pacotes por protocolo de aplica\u00e7\u00e3o"))
    print("  Grafico 4 inserido")

# 8.6 Analise Estatistica
elements.append(make_paragraph(u"8.6 An\u00e1lise Estat\u00edstica", font_size=14, bold=True))
elements.append(make_paragraph(
    u"A an\u00e1lise estat\u00edstica dos dados capturados permite destacar as seguintes observa\u00e7\u00f5es:"
))

observacoes = [
    u"Predomin\u00e2ncia do TCP: Com 57,1% dos pacotes, o TCP confirma-se como o principal protocolo de transporte utilizado nas comunica\u00e7\u00f5es, resultado esperado dado que a maioria dos servi\u00e7os web opera sobre conex\u00f5es TCP.",
    u"Tr\u00e1fego criptografado: O TLS representou 52,3% do volume total de bytes, demonstrando que a grande maioria do tr\u00e1fego capturado utiliza criptografia, refletindo a ado\u00e7\u00e3o generalizada do HTTPS nos servi\u00e7os web atuais.",
    u"Presen\u00e7a significativa do QUIC: Com 442 pacotes (19,4% do total), o protocolo QUIC demonstra sua crescente ado\u00e7\u00e3o como alternativa ao TCP+TLS para comunica\u00e7\u00f5es web, especialmente em servi\u00e7os do Google e navegadores modernos.",
    u"Predomin\u00e2ncia do IPv6: 74,7% dos pacotes utilizaram IPv6, indicando que a rede analisada j\u00e1 opera predominantemente com o protocolo de internet mais recente.",
    u"Tr\u00e1fego DNS: Os 88 pacotes DNS capturados representam as consultas de resolu\u00e7\u00e3o de nomes realizadas durante a navega\u00e7\u00e3o, confirmando o funcionamento adequado do servi\u00e7o DNS na rede.",
    u"Protocolos de controle: A presen\u00e7a de ICMPv6 (230 pacotes) e IGMP (44 pacotes) indica o funcionamento normal dos mecanismos de descoberta de vizinhos e gerenciamento de grupos multicast na rede.",
]

for obs in observacoes:
    p = doc.add_paragraph()
    run = p.add_run(u"\u2022 " + obs)
    run.font.size = Pt(12)
    elements.append(p)

elements.append(make_paragraph(""))

elements.append(make_paragraph(
    u"Os dados estat\u00edsticos coletados demonstram um perfil de tr\u00e1fego t\u00edpico de navega\u00e7\u00e3o "
    u"web moderna, com predomin\u00e2ncia de conex\u00f5es seguras (TLS/HTTPS) e presen\u00e7a significativa do "
    u"protocolo QUIC, refletindo as tend\u00eancias atuais de desenvolvimento e otimiza\u00e7\u00e3o dos protocolos "
    u"de comunica\u00e7\u00e3o na Internet."
))

# Inserir todos os elementos na ordem correta (de tras para frente)
for elem in reversed(elements):
    insert_after(ref, elem)
    
print(f"\nTotal de {len(elements)} elementos inseridos na secao 8")

# ============================================================
# Salvar
# ============================================================
doc.save(DOCX_PATH)
print(f"Documento salvo: {DOCX_PATH}")
print("Secao 8 preenchida com sucesso!")
